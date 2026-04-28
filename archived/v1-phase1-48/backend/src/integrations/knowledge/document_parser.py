"""Document Parser — multi-format document ingestion.

Supports PDF, Word (docx), HTML, Markdown, and plain text.
Falls back to plain text extraction when specialised parsers are unavailable.

Sprint 118 — Phase 38 E2E Assembly C.
"""

import logging
import os
from dataclasses import dataclass, field
from enum import Enum
from typing import List, Optional

logger = logging.getLogger(__name__)


class DocumentFormat(str, Enum):
    """Supported document formats."""

    PDF = "pdf"
    DOCX = "docx"
    HTML = "html"
    MARKDOWN = "markdown"
    TEXT = "text"
    UNKNOWN = "unknown"


@dataclass
class ParsedDocument:
    """Result of document parsing."""

    content: str
    format: DocumentFormat
    title: str = ""
    metadata: dict = field(default_factory=dict)
    page_count: int = 0
    char_count: int = 0


class DocumentParser:
    """Multi-format document parser.

    Detects format from file extension and delegates to the appropriate
    parser.  Falls back to plain text when libraries are not installed.
    """

    FORMAT_MAP = {
        ".pdf": DocumentFormat.PDF,
        ".docx": DocumentFormat.DOCX,
        ".doc": DocumentFormat.DOCX,
        ".html": DocumentFormat.HTML,
        ".htm": DocumentFormat.HTML,
        ".md": DocumentFormat.MARKDOWN,
        ".markdown": DocumentFormat.MARKDOWN,
        ".txt": DocumentFormat.TEXT,
        ".text": DocumentFormat.TEXT,
        ".csv": DocumentFormat.TEXT,
        ".json": DocumentFormat.TEXT,
        ".yaml": DocumentFormat.TEXT,
        ".yml": DocumentFormat.TEXT,
    }

    def detect_format(self, file_path: str) -> DocumentFormat:
        """Detect document format from file extension."""
        ext = os.path.splitext(file_path)[1].lower()
        return self.FORMAT_MAP.get(ext, DocumentFormat.UNKNOWN)

    async def parse(self, file_path: str, content: Optional[bytes] = None) -> ParsedDocument:
        """Parse a document from file path or raw bytes.

        Args:
            file_path: Path to the document file.
            content: Optional raw bytes (if already loaded).

        Returns:
            ParsedDocument with extracted text.
        """
        doc_format = self.detect_format(file_path)
        title = os.path.basename(file_path)

        if content is None:
            try:
                with open(file_path, "rb") as f:
                    content = f.read()
            except FileNotFoundError:
                return ParsedDocument(
                    content="",
                    format=doc_format,
                    title=title,
                    metadata={"error": f"File not found: {file_path}"},
                )

        if doc_format == DocumentFormat.PDF:
            text = self._parse_pdf(content)
        elif doc_format == DocumentFormat.DOCX:
            text = self._parse_docx(content)
        elif doc_format == DocumentFormat.HTML:
            text = self._parse_html(content)
        elif doc_format in (DocumentFormat.MARKDOWN, DocumentFormat.TEXT):
            text = content.decode("utf-8", errors="replace")
        else:
            text = content.decode("utf-8", errors="replace")

        return ParsedDocument(
            content=text,
            format=doc_format,
            title=title,
            metadata={"source": file_path, "format": doc_format.value},
            char_count=len(text),
        )

    async def parse_text(self, text: str, title: str = "inline") -> ParsedDocument:
        """Parse raw text content directly."""
        return ParsedDocument(
            content=text,
            format=DocumentFormat.TEXT,
            title=title,
            char_count=len(text),
        )

    @staticmethod
    def _parse_pdf(content: bytes) -> str:
        """Extract text from PDF bytes."""
        try:
            import io
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(content))
            pages = [page.extract_text() or "" for page in reader.pages]
            return "\n\n".join(pages)
        except ImportError:
            logger.warning("PyPDF2 not installed, falling back to raw decode")
            return content.decode("utf-8", errors="replace")
        except Exception as e:
            logger.error("PDF parsing failed: %s", e)
            return content.decode("utf-8", errors="replace")

    @staticmethod
    def _parse_docx(content: bytes) -> str:
        """Extract text from DOCX bytes."""
        try:
            import io
            from docx import Document
            doc = Document(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs)
        except ImportError:
            logger.warning("python-docx not installed, falling back to raw decode")
            return content.decode("utf-8", errors="replace")
        except Exception as e:
            logger.error("DOCX parsing failed: %s", e)
            return content.decode("utf-8", errors="replace")

    @staticmethod
    def _parse_html(content: bytes) -> str:
        """Extract text from HTML bytes."""
        try:
            from html.parser import HTMLParser
            import io

            class TextExtractor(HTMLParser):
                def __init__(self):
                    super().__init__()
                    self.result: List[str] = []
                    self._skip = False

                def handle_starttag(self, tag, attrs):
                    if tag in ("script", "style"):
                        self._skip = True

                def handle_endtag(self, tag):
                    if tag in ("script", "style"):
                        self._skip = False

                def handle_data(self, data):
                    if not self._skip and data.strip():
                        self.result.append(data.strip())

            html_text = content.decode("utf-8", errors="replace")
            extractor = TextExtractor()
            extractor.feed(html_text)
            return "\n".join(extractor.result)
        except Exception as e:
            logger.error("HTML parsing failed: %s", e)
            return content.decode("utf-8", errors="replace")
